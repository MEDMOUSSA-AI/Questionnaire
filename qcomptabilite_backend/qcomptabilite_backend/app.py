"""
Backend لاستبيان: أثر المحاسبة الدولية على الأسواق المالية
- يسجل كل إجابة في قاعدة بيانات SQLite
- لوحة إدارة (HTML/CSS) لعرض الإجابات
- تصدير أي إجابة، أو كل الإجابات، إلى ملف Word بضغطة واحدة
- رسوم بيانية لإحصائيات الإجابات (لوحة الإدارة + تصدير Word)
"""

import os
import io
import json
import math
import sqlite3
from collections import Counter
from datetime import datetime
from functools import wraps

import matplotlib
matplotlib.use("Agg")  # لا حاجة لواجهة رسومية على الخادم
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import arabic_reshaper
from bidi.algorithm import get_display

from flask import (
    Flask, request, jsonify, render_template, redirect,
    url_for, session, send_file, flash, g
)
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --------------------------------------------------------------------------
# الإعدادات
# --------------------------------------------------------------------------
app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret-change-me")
CORS(app)  # يسمح لصفحة الاستبيان (مستضافة في مكان آخر) بإرسال الإجابات

DB_PATH = os.environ.get("DB_PATH", os.path.join(os.path.dirname(__file__), "responses.db"))

ADMIN_USERNAME = os.environ.get("ADMIN_USERNAME", "admin")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "changeme123")  # غيّرها في بيئة الإنتاج

# --------------------------------------------------------------------------
# خط عربي لاستخدامه في الرسوم البيانية (matplotlib لا يدعم تشكيل الحروف العربية
# تلقائيًا، لذلك نُهيّئ النص عبر arabic_reshaper + python-bidi ثم نرسمه بهذا الخط)
# --------------------------------------------------------------------------
_FONT_PATH = os.path.join(os.path.dirname(__file__), "fonts", "NotoNaskhArabic-Regular.ttf")
if os.path.exists(_FONT_PATH):
    fm.fontManager.addfont(_FONT_PATH)
    AR_FONT = fm.FontProperties(fname=_FONT_PATH)
else:
    AR_FONT = fm.FontProperties()  # احتياطي إن لم يُرفَع ملف الخط بعد


def ar_text(text):
    """يهيّئ نصًا عربيًا (يربط الحروف ويصحّح اتجاه العرض) ليظهر سليمًا في matplotlib."""
    try:
        return get_display(arabic_reshaper.reshape(str(text)))
    except Exception:
        return str(text)


_CHART_COLORS = [
    "#16233D", "#9C7A32", "#1F6F5C", "#C9A961", "#7280A0",
    "#B23A48", "#3B6E8F", "#6B4226", "#4C5B7A", "#8E6C88",
]

# وصف الأسئلة (يُستخدم في لوحة الإدارة وملفات Word)
# opts: قائمة الخيارات الرسمية لكل سؤال كما عُرّفت في نموذج الاستبيان (index.html)
#       — تُستخدم لعرض كل الخيارات دائمًا في الرسوم البيانية، حتى لو حصل خيار على 0 إجابة.
#       الأسئلة بدون "opts" (q18 فقط) هي أسئلة نص حر مفتوحة.
_SCALE_AGREE = [
    {"v": "tout_a_fait_dacc", "ar": "موافق بشدة"},
    {"v": "dacc", "ar": "موافق"},
    {"v": "neutre", "ar": "محايد"},
    {"v": "pas_dacc", "ar": "غير موافق"},
    {"v": "pas_du_tout_dacc", "ar": "غير موافق بشدة"},
]
_SCALE_INFLUENCE = [
    {"v": "tres_fortement", "ar": "بشدة جدًا"},
    {"v": "fortement", "ar": "بشدة"},
    {"v": "moyennement", "ar": "بشكل متوسط"},
    {"v": "faiblement", "ar": "بشكل ضعيف"},
    {"v": "pas_du_tout", "ar": "لا تؤثر إطلاقًا"},
]
_OUI_NON = [{"v": "oui", "ar": "نعم"}, {"v": "non", "ar": "لا"}]
_OUI_NON_SANS_OPINION = _OUI_NON + [{"v": "sans_opinion", "ar": "بدون رأي"}]
_OUI_NON_PEUT_ETRE = _OUI_NON + [{"v": "peut_etre", "ar": "ربما"}]

QUESTIONS = [
    {"id": "q1",  "ar": "الجنس",                                              "fr": "Sexe",
     "opts": [{"v": "homme", "ar": "ذكر"}, {"v": "femme", "ar": "أنثى"}]},
    {"id": "q2",  "ar": "العمر",                                              "fr": "Âge",
     "opts": [{"v": "lt25", "ar": "أقل من 25 سنة"}, {"v": "25_35", "ar": "25–35 سنة"},
              {"v": "36_45", "ar": "36–45 سنة"}, {"v": "gt45", "ar": "أكثر من 45 سنة"}]},
    {"id": "q3",  "ar": "المهنة",                                             "fr": "Profession",
     "opts": [{"v": "etudiant", "ar": "طالب"}, {"v": "comptable", "ar": "محاسب"},
              {"v": "auditeur", "ar": "مدقق حسابات"}, {"v": "investisseur", "ar": "مستثمر"},
              {"v": "autre", "ar": "أخرى"}]},
    {"id": "q4",  "ar": "المستوى الدراسي",                                    "fr": "Niveau d'études",
     "opts": [{"v": "licence", "ar": "إجازة (ليسانس)"}, {"v": "master", "ar": "ماستر"},
              {"v": "doctorat", "ar": "دكتوراه"}, {"v": "autre", "ar": "أخرى"}]},
    {"id": "q5",  "ar": "هل تعرف معايير IFRS؟",                                "fr": "Connaissez-vous les IFRS ?",
     "opts": _OUI_NON},
    {"id": "q6",  "ar": "تطبيق IFRS يحسّن جودة المعلومة المالية",              "fr": "IFRS améliore la qualité de l'information",
     "opts": _SCALE_AGREE},
    {"id": "q7",  "ar": "المعايير الدولية تسهّل مقارنة القوائم المالية",       "fr": "Facilite la comparaison des états financiers",
     "opts": _SCALE_AGREE},
    {"id": "q8",  "ar": "المحاسبة الدولية تعزز الشفافية المالية",              "fr": "Renforce la transparence financière",
     "opts": _SCALE_AGREE},
    {"id": "q9",  "ar": "جودة المعلومة المحاسبية تؤثر على قرارات المستثمرين",  "fr": "Influence sur les décisions des investisseurs",
     "opts": _SCALE_INFLUENCE},
    {"id": "q10", "ar": "تقليل عدم تماثل المعلومات في الأسواق المالية",        "fr": "Réduction de l'asymétrie d'information",
     "opts": _SCALE_AGREE},
    {"id": "q11", "ar": "اعتماد IFRS يجذب المستثمرين الأجانب",                 "fr": "Attire les investisseurs étrangers",
     "opts": _SCALE_AGREE},
    {"id": "q12", "ar": "المحاسبة الدولية تعزز ثقة المستثمرين",                "fr": "Renforce la confiance des investisseurs",
     "opts": _SCALE_AGREE},
    {"id": "q13", "ar": "تطبيق المعايير الدولية يفضّل تطور الأسواق المالية",   "fr": "Favorise le développement des marchés",
     "opts": _SCALE_AGREE},
    {"id": "q14", "ar": "الأسواق أكثر كفاءة مع الشفافية المالية",              "fr": "Marchés plus efficaces avec transparence",
     "opts": _SCALE_AGREE},
    {"id": "q15", "ar": "المحاسبة الدولية تسهّل الوصول إلى التمويل",           "fr": "Facilite l'accès au financement",
     "opts": _SCALE_AGREE},
    {"id": "q16", "ar": "أهم فائدة للمحاسبة الدولية",                          "fr": "Principal avantage",
     "opts": [{"v": "transparence", "ar": "الشفافية"}, {"v": "comparabilite", "ar": "قابلية المقارنة"},
              {"v": "attractivite", "ar": "جاذبية للمستثمرين"}, {"v": "reduction_risques", "ar": "تقليل المخاطر"},
              {"v": "autre", "ar": "أخرى"}]},
    {"id": "q17", "ar": "التوصية باعتماد المعايير الدولية لكل الشركات المدرجة", "fr": "Recommandation d'adoption générale",
     "opts": _OUI_NON},
    {"id": "q18", "ar": "ملاحظات واقتراحات",                                  "fr": "Suggestions et commentaires"},
    # q18 بدون "opts": سؤال نص حر مفتوح
]
QUESTION_MAP = {q["id"]: q for q in QUESTIONS}


# --------------------------------------------------------------------------
# قاعدة البيانات
# --------------------------------------------------------------------------
def get_db():
    if "db" not in g:
        g.db = sqlite3.connect(DB_PATH)
        g.db.row_factory = sqlite3.Row
    return g.db


@app.teardown_appcontext
def close_db(exception=None):
    db = g.pop("db", None)
    if db is not None:
        db.close()


def init_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS responses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at TEXT NOT NULL,
            lang TEXT NOT NULL,
            answers_json TEXT NOT NULL,
            user_agent TEXT,
            ip_address TEXT
        )
    """)
    conn.commit()
    conn.close()


init_db()


# --------------------------------------------------------------------------
# مساعد: تسجيل الدخول
# --------------------------------------------------------------------------
def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not session.get("logged_in"):
            return redirect(url_for("admin_login", next=request.path))
        return view(*args, **kwargs)
    return wrapped


# --------------------------------------------------------------------------
# الصفحة الرئيسية: صفحة الاستبيان نفسها
# --------------------------------------------------------------------------
@app.route("/")
def index():
    return render_template("index.html")


# --------------------------------------------------------------------------
# API: استقبال إجابات الاستبيان
# --------------------------------------------------------------------------
@app.route("/api/submit", methods=["POST"])
def api_submit():
    data = request.get_json(silent=True)
    if not data or "answers" not in data:
        return jsonify({"status": "error", "message": "بيانات غير صالحة"}), 400

    lang = data.get("lang", "ar")
    answers = data["answers"]

    db = get_db()
    db.execute(
        "INSERT INTO responses (created_at, lang, answers_json, user_agent, ip_address) VALUES (?, ?, ?, ?, ?)",
        (
            datetime.utcnow().isoformat(),
            lang,
            json.dumps(answers, ensure_ascii=False),
            request.headers.get("User-Agent", ""),
            request.headers.get("X-Forwarded-For", request.remote_addr or ""),
        ),
    )
    db.commit()
    return jsonify({"status": "ok"})


@app.route("/api/health")
def api_health():
    return jsonify({"status": "ok", "time": datetime.utcnow().isoformat()})


# --------------------------------------------------------------------------
# تسجيل الدخول للوحة الإدارة
# --------------------------------------------------------------------------
@app.route("/admin/login", methods=["GET", "POST"])
def admin_login():
    if request.method == "POST":
        username = request.form.get("username", "")
        password = request.form.get("password", "")
        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            session["logged_in"] = True
            next_url = request.args.get("next") or url_for("admin_dashboard")
            return redirect(next_url)
        flash("اسم المستخدم أو كلمة المرور غير صحيحة")
    return render_template("login.html")


@app.route("/admin/logout")
def admin_logout():
    session.clear()
    return redirect(url_for("admin_login"))


# --------------------------------------------------------------------------
# لوحة الإدارة
# --------------------------------------------------------------------------
@app.route("/admin")
@login_required
def admin_dashboard():
    db = get_db()
    rows = db.execute("SELECT * FROM responses ORDER BY id DESC").fetchall()

    responses = []
    for r in rows:
        answers = json.loads(r["answers_json"])
        responses.append({
            "id": r["id"],
            "created_at": r["created_at"],
            "lang": r["lang"],
            "profession": answers.get("q3", ""),
            "sexe": answers.get("q1", ""),
            "age": answers.get("q2", ""),
        })

    return render_template("dashboard.html", responses=responses, total=len(responses))


# --------------------------------------------------------------------------
# إحصائيات ورسوم بيانية
# --------------------------------------------------------------------------
def compute_question_stats():
    """يحسب توزيع الإجابات لكل سؤال.
    الأسئلة ذات الخيارات المغلقة (opts) تعرض كل خياراتها الرسمية دائمًا — حتى
    لو حصل خيار على 0 إجابة — بعد ترجمة الرموز التقنية المخزّنة (مثل
    tout_a_fait_dacc) إلى تسمياتها العربية الصحيحة (موافق بشدة).
    الأسئلة المفتوحة (بدون opts، أي q18) تُعدّ فيها التكرارات الفعلية للنصوص.
    نوع الرسم: دائري لسؤال بخيارين فقط، أعمدة لما عدا ذلك."""
    db = get_db()
    rows = db.execute("SELECT answers_json FROM responses").fetchall()

    stats = []
    for q in QUESTIONS:
        opts = q.get("opts")

        if opts:
            counts_map = {opt["ar"]: 0 for opt in opts}
            value_to_label = {opt["v"]: opt["ar"] for opt in opts}
            no_answer = 0
            legacy = 0  # إجابات محفوظة بقيم من نسخة سابقة للاستبيان لم تعد ضمن الخيارات الحالية

            for row in rows:
                answers = json.loads(row["answers_json"])
                raw_value = answers.get(q["id"], "")
                if raw_value in value_to_label:
                    counts_map[value_to_label[raw_value]] += 1
                elif raw_value and "أخرى" in counts_map:
                    counts_map["أخرى"] += 1
                elif raw_value:
                    legacy += 1
                else:
                    no_answer += 1

            labels = list(counts_map.keys())
            counts = list(counts_map.values())
            if no_answer:
                labels.append("بدون إجابة")
                counts.append(no_answer)
            if legacy:
                labels.append("قيم سابقة (سلّم قديم)")
                counts.append(legacy)
        else:
            counter = Counter()
            for row in rows:
                answers = json.loads(row["answers_json"])
                value = str(answers.get(q["id"], "") or "").strip()
                counter[value if value else "بدون إجابة"] += 1
            labels = list(counter.keys())
            counts = list(counter.values())

        total_q = sum(counts)
        stats.append({
            "id": q["id"],
            "label": q["ar"],
            "label_fr": q["fr"],
            "chart_type": "pie" if len(labels) == 2 else "bar",
            "labels": labels,
            "counts": counts,
            "total": total_q,
        })
    return stats


def generate_stat_chart_image(stat):
    """يولّد صورة PNG (في الذاكرة) لرسم بياني يمثّل توزيع إجابات سؤال واحد —
    دائري للأسئلة ذات خيارين، وأعمدة أفقية لما عداها — بنفس ألوان لوحة الإدارة."""
    labels = stat["labels"]
    counts = stat["counts"]
    total = sum(counts) or 1
    colors = [_CHART_COLORS[i % len(_CHART_COLORS)] for i in range(len(labels))]

    fig, ax = plt.subplots(figsize=(7.4, 4.3))

    if stat["chart_type"] == "pie":
        wedges, _ = ax.pie(
            counts, colors=colors, startangle=90, counterclock=False,
            wedgeprops={"edgecolor": "#FCFBF8", "linewidth": 2},
        )
        for w, c in zip(wedges, counts):
            if c == 0:
                continue
            ang = math.radians((w.theta2 + w.theta1) / 2)
            x, y = 0.68 * math.cos(ang), 0.68 * math.sin(ang)
            pct = f"{c / total * 100:.1f}%".replace(".", ",")
            ax.text(x, y, ar_text(pct), ha="center", va="center", color="white",
                     fontproperties=AR_FONT, fontsize=11)
        ax.legend(
            wedges, [ar_text(l) for l in labels], loc="center left",
            bbox_to_anchor=(1.02, 0.5), prop=AR_FONT, frameon=False, fontsize=10.5,
        )
        ax.axis("equal")
    else:
        y_pos = list(range(len(labels)))
        bars = ax.barh(y_pos, counts, color=colors, height=0.6)
        ax.set_yticks(y_pos)
        ax.set_yticklabels([ar_text(l) for l in labels], fontproperties=AR_FONT, fontsize=10.5)
        ax.invert_yaxis()
        for spine in ("top", "right", "left"):
            ax.spines[spine].set_visible(False)
        ax.get_xaxis().set_visible(False)
        max_count = max(counts) if counts else 1
        for bar, c in zip(bars, counts):
            pct = f"{c / total * 100:.1f}%".replace(".", ",")
            ax.text(
                bar.get_width() + max_count * 0.02, bar.get_y() + bar.get_height() / 2,
                ar_text(f"{c} ({pct})"), va="center", ha="left",
                fontproperties=AR_FONT, fontsize=10, color="#16233D",
            )

    ax.set_title(ar_text(stat["label"]) + "\n", fontproperties=AR_FONT, fontsize=13.5,
                 fontweight="bold", color="#16233D", loc="center")
    fig.text(0.5, 0.93, ar_text(f"{total} إجابة"), ha="center",
              fontproperties=AR_FONT, fontsize=10, color="#9C7A32")

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=170, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


@app.route("/admin/export-stats/word")
@login_required
def admin_export_stats_word():
    stats = compute_question_stats()
    db = get_db()
    total = db.execute("SELECT COUNT(*) AS c FROM responses").fetchone()["c"]

    doc = Document()
    style_document_base(doc)

    title = doc.add_heading(level=0)
    trun = title.add_run("استبيان: أثر المحاسبة الدولية على الأسواق المالية — الإحصائيات")
    set_rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = doc.add_paragraph()
    summary_run = summary.add_run(
        f"عدد الإجابات: {total}   |   تاريخ التصدير: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}"
    )
    set_rtl(summary)
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for stat in stats:
        img_buf = generate_stat_chart_image(stat)
        doc.add_picture(img_buf, width=Cm(16))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"statistiques_{datetime.utcnow().strftime('%Y%m%d_%H%M')}.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/admin/stats")
@login_required
def admin_stats():
    stats = compute_question_stats()
    db = get_db()
    total = db.execute("SELECT COUNT(*) AS c FROM responses").fetchone()["c"]
    return render_template("stats.html", stats=stats, total=total)


@app.route("/admin/stats/chart/<question_id>.png")
@login_required
def admin_stat_chart_image(question_id):
    stats = compute_question_stats()
    stat = next((s for s in stats if s["id"] == question_id), None)
    if stat is None:
        flash("سؤال غير معروف")
        return redirect(url_for("admin_stats"))
    img_buf = generate_stat_chart_image(stat)
    return send_file(img_buf, mimetype="image/png")


@app.route("/admin/response/<int:response_id>")
@login_required
def admin_response_detail(response_id):
    db = get_db()
    row = db.execute("SELECT * FROM responses WHERE id = ?", (response_id,)).fetchone()
    if row is None:
        flash("لم يتم العثور على هذه الإجابة")
        return redirect(url_for("admin_dashboard"))

    answers = json.loads(row["answers_json"])
    items = []
    for q in QUESTIONS:
        items.append({
            "label": q["ar"],
            "label_fr": q["fr"],
            "value": answers.get(q["id"], ""),
        })

    return render_template("detail.html", response=row, items=items)


@app.route("/admin/response/<int:response_id>/delete", methods=["POST"])
@login_required
def admin_response_delete(response_id):
    db = get_db()
    db.execute("DELETE FROM responses WHERE id = ?", (response_id,))
    db.commit()
    return redirect(url_for("admin_dashboard"))


# --------------------------------------------------------------------------
# تصدير Word
# --------------------------------------------------------------------------
def set_rtl(paragraph):
    """يجعل الفقرة من اليمين لليسار (لدعم النص العربي في Word)."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def set_cell_rtl(cell):
    for p in cell.paragraphs:
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def style_document_base(doc):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    # يضمن أن الخط العربي يُعرض بشكل صحيح
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:cs"), "Arial")


def add_response_to_doc(doc, row, answers, index=None):
    ink = RGBColor(0x16, 0x23, 0x3D)
    bronze = RGBColor(0x9C, 0x7A, 0x32)

    heading_text = f"الإجابة رقم {row['id']}" if index is None else f"الإجابة رقم {row['id']} ({index})"
    h = doc.add_heading(level=1)
    run = h.add_run(heading_text)
    run.font.color.rgb = ink
    set_rtl(h)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"تاريخ الإرسال: {row['created_at']}  |  اللغة: {row['lang']}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = bronze
    set_rtl(meta)
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "السؤال"
    hdr[1].text = "الإجابة"
    set_cell_rtl(hdr[0])
    set_cell_rtl(hdr[1])

    for q in QUESTIONS:
        value = answers.get(q["id"], "")
        if q["id"] + "_autre" in answers and answers.get(q["id"]) == "autre":
            value = f"أخرى: {answers.get(q['id'] + '_autre', '')}"
        row_cells = table.add_row().cells
        row_cells[0].text = q["ar"]
        row_cells[1].text = str(value) if value else "—"
        set_cell_rtl(row_cells[0])
        set_cell_rtl(row_cells[1])

    doc.add_paragraph()


@app.route("/admin/response/<int:response_id>/word")
@login_required
def admin_export_word(response_id):
    db = get_db()
    row = db.execute("SELECT * FROM responses WHERE id = ?", (response_id,)).fetchone()
    if row is None:
        flash("لم يتم العثور على هذه الإجابة")
        return redirect(url_for("admin_dashboard"))

    answers = json.loads(row["answers_json"])

    doc = Document()
    style_document_base(doc)
    title = doc.add_heading(level=0)
    trun = title.add_run("استبيان: أثر المحاسبة الدولية على الأسواق المالية")
    set_rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_response_to_doc(doc, row, answers)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"response_{response_id}.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/admin/export-all/word")
@login_required
def admin_export_all_word():
    db = get_db()
    rows = db.execute("SELECT * FROM responses ORDER BY id ASC").fetchall()

    doc = Document()
    style_document_base(doc)
    title = doc.add_heading(level=0)
    title.add_run("استبيان: أثر المحاسبة الدولية على الأسواق المالية — كل الإجابات")
    set_rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = doc.add_paragraph()
    summary_run = summary.add_run(f"عدد الإجابات: {len(rows)}   |   تاريخ التصدير: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}")
    set_rtl(summary)
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, row in enumerate(rows, start=1):
        answers = json.loads(row["answers_json"])
        doc.add_page_break()
        add_response_to_doc(doc, row, answers, index=i)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"all_responses_{datetime.utcnow().strftime('%Y%m%d_%H%M')}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
